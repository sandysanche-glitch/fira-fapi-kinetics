# save as make_method_docx.py and run: python make_method_docx.py
from docx import Document
from docx.shared import Pt

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def H(t, lvl=1): doc.add_heading(t, level=lvl)
def P(t): doc.add_paragraph(t)
def B(t): doc.add_paragraph(t, style='List Bullet')
def N(t): doc.add_paragraph(t, style='List Number')

H("Adapted Method (No Frames; Rank-to-Time Mapping)", 1)
P("Version 1.0 — FAPI dataset focus")

H("Core Idea", 2)
P("Use the area rank of each crystal to define a pseudo-time for nucleation (0–60 ms). "
  "Growth proceeds from that nucleation time to 0–600 ms.")

H("1) Inputs per Object", 2)
for x in [
  "category_id: 1 = crystal (mask), 2 = nucleus (point), 3 = defect (mask/attribute)",
  "Crystal mask → area A_i, perimeter P_i",
  "Nucleation center (x_i, y_i) (if available)",
  "Defect area A_def,i (class 3 or overlap)",
  "Optional: global coordinates (for bulk X(t) only)",
]: B(x)

H("2) Derived Features", 2)
B("Circularity: C_i = 4π A_i / P_i^2 ∈ (0,1]")
B("Defect fraction: φ_i = A_def,i / (A_i + ε), ε≈1e-6")

H("3) Rank-to-Time Mapping", 2)
P("Sort by final area A_i (ascending). q_i = (rank(A_i) − 1) / (N − 1) ∈ [0,1].")
P("Set t0_i = 60 ms × q_i (smallest → ~0 ms; largest → ~60 ms).")

H("4) Growth (0–600 ms) with Morphology Penalties", 2)
P("R_i = sqrt(A_i/π); Δt_i = 600 − t0_i (ms).")
for x in [
  "f_circ,i = exp[−α(1 − C_i)]",
  "f_def,i  = exp[−β φ_i]",
  "v_i = v0 · f_circ,i · f_def,i",
  "Choose v0: R_i = v_i Δt_i ⇒ v0 = median_i { R_i / (Δt_i f_circ,i f_def,i) }",
  "r_i(t)=v_i·max(0,t−t0_i); A_pred,i(t)=π r_i(t)^2, t∈[0,600] ms",
]: B(x)

H("5) Nucleation Rate Density I(t)", 2)
for x in [
  "Fit on t ∈ (0,60] ms using t0_i:",
  "Lognormal: I(t) = [1/(t σ √(2π))] · exp( −(ln t − μ)^2 / (2σ^2) )",
  "Gamma: I(t) = [1/(Γ(k) θ^k)] · t^{k−1} · exp( −t/θ )",
  "Select via BIC (AIC tiebreaker). Bootstrap t0_i for 95% bands. 1-ms grid."
]: B(x)

H("6) Bulk Fraction X(t) (Optional)", 2)
P("Estimate A_eff (e.g., convex hull area). X_pred(t) ≈ (1/A_eff) Σ_i A_pred,i(t).")
P("Overlay Avrami: X_Avrami(t) = 1 − exp[−K t^n] (fix n, fit K).")

H("7) Outputs", 2)
for x in [
  "CSVs: I(t); per-object A_pred,i(t); optional X(t)+Avrami; parameters; AIC/BIC.",
  "Plots: I(t) with selected model; per-object growth overlays; optional X(t) vs Avrami."
]: B(x)

H("Practical Procedure (Step-by-Step)", 2)
for s in [
  "Load COCO JSONs; extract crystals (1) and defects (3).",
  "Compute A_i, P_i, C_i; derive φ_i (or 0).",
  "Sort by A_i; compute q_i and t0_i = 60 ms × q_i.",
  "Compute f_circ,i, f_def,i (α=β=1 to start).",
  "Compute R_i = sqrt(A_i/π) and Δt_i = 600 − t0_i.",
  "Set v0 = median_i { R_i / (Δt_i f_circ,i f_def,i) }.",
  "Generate A_pred,i(t) on 1-ms grid; cap at observed A_i at 600 ms.",
  "Fit I(t) (lognormal vs gamma); choose via BIC; bootstrap bands.",
  "Optional: A_eff → X_pred(t); fit Avrami K for fixed n.",
  "Export CSVs and plots."
]: N(s)

H("Defaults, Tuning, Notes", 2)
for x in [
  "Windows: nucleation 0–60 ms; growth 0–600 ms; grid 1 ms.",
  "Penalty weights α, β ≥ 0 (default 1.0).",
  "If A_eff is uncertain, skip bulk X(t); use per-object diagnostics.",
  "Rank-to-time is robust with one snapshot per object.",
]: B(x)

doc.save("Adapted_Method_rank_to_time_FAPI.docx")
print("Wrote Adapted_Method_rank_to_time_FAPI.docx")
